import docx
from docx.shared import Pt
import re

def save_solution_to_doc(solution_data, file_path, is_digital_one=False):
    """
    Generates a .docx document from the agent's solution data.
    Supports both Ensemble (Code) and DigitalOne (Report) formats.
    """
    try:
        doc = docx.Document()
        
        # --- TITLE ---
        if is_digital_one:
            doc.add_heading('Digital One Product Specification', level=0)
        else:
            doc.add_heading('Ensemble Technical Design', level=0)
        
        # --- 1. SUMMARY (Common) ---
        doc.add_heading('1. Executive Summary', level=1)
        doc.add_paragraph(solution_data.get('summary', 'Not provided.'))

        # --- 2. ASSUMPTIONS (Common) ---
        if solution_data.get("assumptions"):
            doc.add_heading('2. Assumptions & Prerequisites', level=1)
            for assumption in solution_data["assumptions"]:
                doc.add_paragraph(assumption, style='List Bullet')

        # --- 3. DIGITAL ONE REPORT (Special Logic) ---
        if is_digital_one:
            # D1 workflow puts the main content in 'markdown_report' or just uses summary
            report_content = solution_data.get("markdown_report", "")
            if report_content:
                doc.add_heading('3. Detailed Analysis', level=1)
                # Basic Markdown cleaning for Word
                for line in report_content.split('\n'):
                    # Handle basic headers
                    if line.startswith('## '):
                        doc.add_heading(line.replace('## ', ''), level=2)
                    elif line.startswith('### '):
                        doc.add_heading(line.replace('### ', ''), level=3)
                    elif line.strip().startswith('|'):
                        # Table row - simplified handling (just print as text for now)
                        p = doc.add_paragraph()
                        p.add_run(line).font.name = 'Courier New'
                    else:
                        doc.add_paragraph(line)
            
            doc.save(file_path)
            print(f"✅ Successfully saved D1 Spec to: {file_path}")
            return

        # --- 3. ENSEMBLE IMPACT & CODE (Default Logic) ---
        
        doc.add_heading('3. Impact Analysis', level=1)
        doc.add_paragraph(solution_data.get('impact_analysis', 'Not provided.'))

        if solution_data.get("user_stories"):
            doc.add_heading('4. Proposed User Stories', level=1)
            for story in solution_data["user_stories"]:
                p = doc.add_paragraph(style='List Bullet')
                title = story.get('title') or story.get('summary') or 'Story'
                p.add_run(title).bold = True
                
                ac = story.get('acceptance_criteria', [])
                if isinstance(ac, list):
                    for criteria in ac:
                        doc.add_paragraph(f"  - {criteria}", style='List 2')
                else:
                    doc.add_paragraph(f"  - {ac}", style='List 2')

        if solution_data.get("code_changes"):
            doc.add_heading('5. Proposed Code Changes', level=1)
            for change in solution_data["code_changes"]:
                doc.add_heading(f"File: {change.get('file_path', 'Unknown file')}", level=3)
                p = doc.add_paragraph()
                # Use monospaced font for code diffs
                run = p.add_run(change.get('diff', 'No diff').replace('\\n', '\n'))
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
        
        if solution_data.get("doc_changes"):
             doc.add_heading('6. Documentation Updates', level=1)
             for change in solution_data["doc_changes"]:
                doc.add_heading(f"File: {change.get('file_path', 'Unknown file')}", level=3)
                p = doc.add_paragraph()
                run = p.add_run(change.get('diff', 'No diff').replace('\\n', '\n'))
                run.font.name = 'Courier New'
                run.font.size = Pt(10)

        doc.save(file_path)
        print(f"✅ Successfully saved Impact Analysis to: {file_path}")
        
    except Exception as e:
        print(f"🚨 Error saving .docx file: {e}")